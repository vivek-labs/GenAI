from pathlib import Path
from docx import Document as DocxDocument
from src.ingestion.extractor_factory import ExtractorFactory
from src.chunking.sentence_chunker import SentenceChunker
from src.ingestion.docx_extractor import DOCXExtractor
def test_docx_ingestion_and_chunking_end2end(tmp_path: Path):
    #Arrange: Create a valid DOCX test file
    file_path = tmp_path / "rag_notes.docx"

    docx_file = DocxDocument()
    docx_file.add_heading("Retrieval-Augmented Generation", level=1)
    docx_file.add_paragraph(
        "Retrieval-Augmented Generation combines retrieval with a language model."
    )
    docx_file.add_paragraph(
        "Embeddings convert document text into dense numerical vectors."
    )
    docx_file.add_paragraph(
        "ChromaDB stores vectors and performs semantic similarity search."
    )
    docx_file.save(file_path)

    #Act: Factory selects the appropriate extractor
    extractor = ExtractorFactory.create(str(file_path))
    documents = extractor.extract(str(file_path))

    chunker = SentenceChunker(
        chunk_size=120,
        overlap_sentences=1,
    )

    chunks = chunker.chunk(documents)

    #Assert: Factory selected DOCXExtractor
    assert isinstance(extractor, DOCXExtractor)

    # Assert: DOCX extraction worked
    assert len(documents) == 1

    extracted_document = documents[0]

    assert extracted_document.document_id == "rag_notes"
    assert extracted_document.metadata.file_type == "docx"
    assert extracted_document.metadata.page_number == 1
    assert extracted_document.metadata.source_file == str(file_path)

    assert "Retrieval-Augmented Generation" in extracted_document.text
    assert "Embeddings convert document text" in extracted_document.text
    assert "ChromaDB stores vectors" in extracted_document.text

    # Assert: chunking worked
    assert len(chunks) > 0

    combined_chunk_text = " ".join(chunk.text for chunk in chunks)


    assert "Retrieval-Augmented Generation" in combined_chunk_text
    assert "Embeddings convert document text" in combined_chunk_text
    assert "ChromaDB stores vectors" in combined_chunk_text

     # Assert: metadata continues from Document to Chunk
    for chunk in chunks:
        assert chunk.metadata.file_type == "docx"
        assert chunk.metadata.source_file == str(file_path)



    